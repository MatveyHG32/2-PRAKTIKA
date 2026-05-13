"""Generate the Practice 3 final report (.docx).

Outputs: docs/reports/Отчёт_практика_3.docx

Структура (по шаблону из приложенного образца):
титул → индивидуальное задание → введение → 1. Проектирование архитектуры
приложения → 2. Основные компоненты и их функции → 3. Алгоритм работы
приложения → 4. Индивидуальное задание → Заключение → Приложение.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "reports"
OUT_DIR.mkdir(parents=True, exist_ok=True)

SCREENS = ROOT / "docs" / "screenshots"
FIGMA = ROOT / "docs" / "figma_mockups"


def add_centered(doc, text, *, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_paragraph(doc, text, *, size=14, indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_heading(doc, text, *, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_image(doc, path: Path, caption: str, width_inches: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(12)
    cap_run.font.name = "Times New Roman"


def add_code(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(11)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    # ===== Титул =====
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", size=12)
    add_centered(doc, "ФГБОУ ВО «Брянский государственный инженерно-технологический университет»", size=12)
    add_centered(doc, "Многопрофильный колледж", size=12)
    add_centered(doc, "Кафедра «Информационные технологии»", size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "ОТЧЁТ", bold=True, size=18)
    add_centered(doc, "по производственной практике", size=14)
    add_centered(doc, "по технологии разработки программного обеспечения", size=14)
    add_centered(doc, "(практика 3)", size=14)
    add_centered(doc, "ОП – 02068025-09.02.07-085.25", size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "Студент _________________________________________ Магасумов М. М.", indent=False)
    add_paragraph(doc, "Группа ИСП(спо)-209-1                          № зачётной книжки: 23-2.085", indent=False)
    add_paragraph(doc, "Руководитель от вуза ________________ преподаватель И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Нормоконтроль ______________________ преподаватель И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Доступ к защите «24» 06 2026 г. ___ преподаватель И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Дата защиты «25» 06 2026 г.  Оценка __________________", indent=False)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Брянск 2026", size=12)
    doc.add_page_break()

    # ===== Индивидуальное задание =====
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", size=12)
    add_centered(doc, "Федеральное государственное бюджетное образовательное учреждение", size=12)
    add_centered(doc, "высшего образования", size=12)
    add_centered(doc, "«Брянский государственный инженерно-технологический университет»", size=12)
    add_centered(doc, "Многопрофильный колледж", size=12)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Индивидуальное задание", bold=True, size=16)
    add_centered(doc, "по производственной практике", size=14)
    doc.add_paragraph()
    add_paragraph(doc, "Обучающийся: Магасумов М. М.", indent=False)
    add_paragraph(doc, "Руководитель практики от вуза: И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Сроки прохождения практики: 08.06.2026 – 25.06.2026", indent=False)
    add_paragraph(doc, "Место прохождения практики: кафедра ИТ БГИТУ", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Содержание индивидуального задания:", bold=True, indent=False)
    add_paragraph(doc, "ВВЕДЕНИЕ ............................................................................................................. 3", indent=False)
    add_paragraph(doc, "1. Проектирование архитектуры приложения ...................................................... 4", indent=False)
    add_paragraph(doc, "2. Основные компоненты и их функции ............................................................... 6", indent=False)
    add_paragraph(doc, "3. Алгоритм работы приложения ........................................................................... 9", indent=False)
    add_paragraph(doc, "4. Индивидуальное задание ..................................................................................11", indent=False)
    add_paragraph(doc, "ЗАКЛЮЧЕНИЕ .....................................................................................................14", indent=False)
    add_paragraph(doc, "ПРИЛОЖЕНИЕ ....................................................................................................15", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Дата выдачи задания 08 июня 2026 г.", indent=False)
    add_paragraph(doc, "Руководитель практики ________________ И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Задание принял к исполнению ____________", indent=False)
    doc.add_page_break()

    # ===== Введение =====
    add_centered(doc, "ВВЕДЕНИЕ", bold=True, size=16)
    add_paragraph(
        doc,
        "Разработка программного обеспечения (англ. software development) — деятельность "
        "по созданию нового программного обеспечения, состоящая из этапов анализа "
        "предметной области, проектирования, кодирования, тестирования и сопровождения.",
    )
    add_paragraph(
        doc,
        "Целью производственной практики 3 является закрепление практических навыков "
        "разработки клиент-серверного веб-приложения и подготовка итоговой "
        "документации к защите учебного проекта «Планировщик задач».",
    )
    add_paragraph(doc, "Задачи производственной практики:", bold=True, indent=False)
    add_paragraph(doc, "— оформить отчётную документацию по итогам трёх практик;", indent=False)
    add_paragraph(doc, "— подготовить два демонстрационных видео по проекту (структура кода и пользовательская демонстрация);", indent=False)
    add_paragraph(doc, "— заполнить дневник производственной практики;", indent=False)
    add_paragraph(doc, "— подготовиться к защите учебного проекта;", indent=False)
    add_paragraph(doc, "— описать архитектуру приложения, его основные компоненты и алгоритм работы;", indent=False)
    add_paragraph(doc, "— провести функциональное тестирование приложения и устранить найденные замечания.", indent=False)
    doc.add_page_break()

    # ===== 1. Проектирование архитектуры приложения =====
    add_centered(doc, "1. Проектирование архитектуры приложения", bold=True, size=16)
    add_paragraph(
        doc,
        "Приложение «Планировщик задач» представляет собой клиент-серверное "
        "веб-приложение, реализованное по архитектурному шаблону MVC "
        "(Model-View-Controller) в его серверной разновидности (server-side MVC) "
        "со следующим распределением ответственности:",
    )
    add_paragraph(doc, "— Модель (Model): SQLAlchemy ORM-классы User, Project, Task, Comment в модуле app/models.py;", indent=False)
    add_paragraph(doc, "— Представление (View): шаблоны Jinja2 в каталоге templates/ (base.html, dashboard.html, task_detail.html, projects.html, login.html, register.html);", indent=False)
    add_paragraph(doc, "— Контроллер (Controller): роутеры FastAPI в каталоге app/routers/ (auth.py, pages.py, tasks.py, projects.py).", indent=False)

    add_heading(doc, "1.1. Используемые технологии")
    add_paragraph(doc, "Серверная часть:", bold=True, indent=False)
    add_paragraph(doc, "— язык программирования Python 3.12;", indent=False)
    add_paragraph(doc, "— веб-фреймворк FastAPI 0.115 (асинхронный, основан на Starlette и Pydantic);", indent=False)
    add_paragraph(doc, "— ORM SQLAlchemy 2.0;", indent=False)
    add_paragraph(doc, "— СУБД SQLite (файл tasks.db в корне проекта);", indent=False)
    add_paragraph(doc, "— шаблонизатор Jinja2 (через fastapi.templating.Jinja2Templates);", indent=False)
    add_paragraph(doc, "— bcrypt — хеширование паролей пользователей;", indent=False)
    add_paragraph(doc, "— Starlette SessionMiddleware — серверные сессии через cookie.", indent=False)
    add_paragraph(doc, "Клиентская часть:", bold=True, indent=False)
    add_paragraph(doc, "— HTML5 (семантическая разметка);", indent=False)
    add_paragraph(doc, "— чистый CSS3 (без фреймворков), CSS Grid и Flexbox, медиа-запросы;", indent=False)
    add_paragraph(doc, "— минимальный нативный JavaScript для обработки модальных окон.", indent=False)
    add_paragraph(doc, "Сервер запуска:", bold=True, indent=False)
    add_paragraph(doc, "— ASGI-сервер Uvicorn 0.30.", indent=False)

    add_heading(doc, "1.2. Структура каталогов проекта")
    add_code(doc,
        "task-planner/\n"
        "├── app/\n"
        "│   ├── database.py        — настройка SQLAlchemy и сессии\n"
        "│   ├── models.py          — модели User, Project, Task, Comment\n"
        "│   ├── security.py        — bcrypt-хеширование паролей\n"
        "│   ├── deps.py            — зависимости (текущий пользователь)\n"
        "│   └── routers/\n"
        "│       ├── auth.py        — регистрация, логин, логаут\n"
        "│       ├── pages.py       — GET-страницы (dashboard, task, projects)\n"
        "│       ├── tasks.py       — CRUD задач и комментариев\n"
        "│       └── projects.py    — CRUD проектов\n"
        "├── templates/             — шаблоны Jinja2\n"
        "├── static/\n"
        "│   ├── css/style.css      — единый файл стилей (~2000 строк)\n"
        "│   └── js/app.js          — обработчики модальных окон\n"
        "├── docs/                  — дизайн-макеты и отчёты\n"
        "├── scripts/               — служебные скрипты\n"
        "├── seed_demo.py           — заливка демонстрационных данных\n"
        "├── main.py                — точка входа\n"
        "└── requirements.txt       — зависимости\n"
    )
    doc.add_page_break()

    # ===== 2. Основные компоненты и их функции =====
    add_centered(doc, "2. Основные компоненты и их функции", bold=True, size=16)

    add_heading(doc, "2.1. Модели данных")
    add_paragraph(doc,
        "В app/models.py определены четыре основные сущности предметной области, "
        "связанные между собой по внешним ключам.")
    add_paragraph(doc, "— User — пользователь системы (логин, отображаемое имя, хеш пароля);", indent=False)
    add_paragraph(doc, "— Project — проект для группировки задач (название, описание, цвет);", indent=False)
    add_paragraph(doc, "— Task — задача (название, описание, статус, приоритет, дедлайн, автор, исполнитель, проект);", indent=False)
    add_paragraph(doc, "— Comment — комментарий к задаче (текст, автор, дата).", indent=False)
    add_paragraph(doc,
        "Статус задачи (todo, in_progress, done) и приоритет (low, medium, high) "
        "реализованы как перечисления (Enum) — это гарантирует корректность данных "
        "на уровне базы.")

    add_heading(doc, "2.2. Аутентификация")
    add_paragraph(doc,
        "За авторизацию отвечает app/routers/auth.py. При регистрации пароль "
        "пользователя хешируется библиотекой bcrypt (12 раундов соли по умолчанию). "
        "При входе сравнивается хеш сохранённого пароля с введённым. Идентификатор "
        "пользователя помещается в подписанный cookie через Starlette SessionMiddleware.")

    add_heading(doc, "2.3. CRUD задач")
    add_paragraph(doc,
        "Роутер app/routers/tasks.py реализует полный CRUD: создание (POST /tasks/create), "
        "обновление полей (POST /tasks/{id}/update), быстрая смена статуса "
        "(POST /tasks/{id}/status), удаление (POST /tasks/{id}/delete) и добавление "
        "комментариев (POST /tasks/{id}/comments). Все POST-формы возвращают 303 "
        "See Other — стандартный паттерн POST → Redirect → GET, чтобы пользователь "
        "после обновления страницы не подтверждал повторную отправку формы.")

    add_heading(doc, "2.4. Дашборд (канбан-доска)")
    add_paragraph(doc,
        "Главная страница приложения (app/routers/pages.py) формирует канбан-доску "
        "из трёх колонок: «К выполнению», «В работе», «Выполнено». Над доской "
        "располагаются пять статистических плиток с подсчётом задач по статусам и "
        "числом просроченных. Каждая карточка задачи показывает цветную полосу "
        "приоритета, чипы проекта и приоритета, описание, дедлайн, аватар "
        "исполнителя и счётчик комментариев.")

    add_heading(doc, "2.5. Адаптивная вёрстка")
    add_paragraph(doc,
        "Стили вынесены в один файл static/css/style.css объёмом около 2000 строк. "
        "Цветовая палитра, отступы и радиусы скругления заданы CSS-переменными в "
        "корневом селекторе :root, что упрощает поддержку. Адаптивная вёрстка "
        "реализована медиа-запросами с точкой перегиба 920 px: доска перестраивается "
        "в одну колонку, статистика — в две, навигация — на новую строку.")

    doc.add_page_break()

    # ===== 3. Алгоритм работы приложения =====
    add_centered(doc, "3. Алгоритм работы приложения", bold=True, size=16)
    add_paragraph(doc,
        "Типичный сценарий взаимодействия пользователя с приложением.")
    add_paragraph(doc, "Шаг 1. Запуск сервера", bold=True, indent=False)
    add_paragraph(doc,
        "При запуске Uvicorn инициализирует приложение FastAPI, создаёт все таблицы "
        "в SQLite (если их ещё нет), подключает middleware сессий и роутеры. После "
        "этого сервер начинает слушать порт 8000.")
    add_paragraph(doc, "Шаг 2. Авторизация", bold=True, indent=False)
    add_paragraph(doc,
        "Пользователь открывает в браузере http://localhost:8000, не имея cookie "
        "сессии. Сервер отвечает редиректом на /login. Пользователь вводит логин и "
        "пароль; форма отправляет POST на /login, сервер проверяет хеш пароля и "
        "при успехе устанавливает cookie сессии с user_id.")
    add_paragraph(doc, "Шаг 3. Просмотр дашборда", bold=True, indent=False)
    add_paragraph(doc,
        "После авторизации сервер обрабатывает GET / — выбирает все задачи "
        "пользователя из БД, считает статистику и передаёт данные в шаблон "
        "dashboard.html. Шаблонизатор Jinja2 рендерит HTML-страницу, которая "
        "и отправляется браузеру.")
    add_paragraph(doc, "Шаг 4. Создание задачи", bold=True, indent=False)
    add_paragraph(doc,
        "Пользователь нажимает «+ Новая задача», открывается модальное окно. После "
        "заполнения формы и отправки браузер делает POST /tasks/create. Контроллер "
        "сохраняет задачу в БД, после чего возвращает редирект на страницу задачи "
        "(GET /tasks/{id}).")
    add_paragraph(doc, "Шаг 5. Редактирование и смена статуса", bold=True, indent=False)
    add_paragraph(doc,
        "На странице задачи доступна форма с полями названия, описания, приоритета, "
        "статуса, срока, проекта и исполнителя. Изменение статуса возможно также "
        "напрямую с карточки на дашборде через выпадающий список. Все изменения "
        "сохраняются в БД и сразу видны в интерфейсе.")
    add_paragraph(doc, "Шаг 6. Комментирование", bold=True, indent=False)
    add_paragraph(doc,
        "В нижней части страницы задачи находится секция «Комментарии». Любой "
        "авторизованный пользователь может оставить комментарий; история "
        "обсуждения остаётся привязанной к задаче.")
    doc.add_page_break()

    # ===== 4. Индивидуальное задание =====
    add_centered(doc, "4. Индивидуальное задание", bold=True, size=16)
    add_paragraph(doc,
        "В рамках индивидуального задания требовалось спроектировать и реализовать "
        "веб-приложение «Планировщик задач» со следующим минимально необходимым "
        "набором функций (приведённым в задании на практику):")
    add_paragraph(doc, "1) добавлять задачи;", indent=False)
    add_paragraph(doc, "2) устанавливать приоритеты задач;", indent=False)
    add_paragraph(doc, "3) назначать задачи другим пользователям;", indent=False)
    add_paragraph(doc, "4) отслеживать выполнение задач.", indent=False)
    add_paragraph(doc,
        "Все четыре функции реализованы и проверены: задачи создаются через "
        "модальную форму, приоритеты выбираются из выпадающего списка "
        "(низкий / средний / высокий), назначение исполнителя выполняется на "
        "странице задачи, а выполнение отслеживается через канбан-доску и "
        "статистические плитки.")
    add_paragraph(doc,
        "Дополнительно сверх задания реализованы: проекты для группировки задач, "
        "комментарии к задачам, прогресс-бары выполнения проектов, адаптивная "
        "вёрстка для мобильных устройств.")

    if (SCREENS / "web_01_dashboard.png").exists():
        add_image(doc, SCREENS / "web_01_dashboard.png",
                  "Рисунок 1 — Главный экран приложения")
    if (SCREENS / "web_02_task_detail.png").exists():
        add_image(doc, SCREENS / "web_02_task_detail.png",
                  "Рисунок 2 — Страница задачи с возможностью редактирования")
    if (SCREENS / "mobile_01_dashboard.png").exists():
        add_image(doc, SCREENS / "mobile_01_dashboard.png",
                  "Рисунок 3 — Мобильная версия дашборда", width_inches=3.0)

    doc.add_page_break()

    # ===== Заключение =====
    add_centered(doc, "ЗАКЛЮЧЕНИЕ", bold=True, size=16)
    add_paragraph(doc,
        "В ходе трёх этапов учебной практики разработано полноценное веб-приложение "
        "«Планировщик задач» — от первоначальной проработки требований и "
        "дизайн-макетов в Figma до работающего серверного приложения на FastAPI "
        "с базой данных SQLite.")
    add_paragraph(doc,
        "Проект развивался последовательно: на первой практике были составлены "
        "функциональные требования и спроектирован дизайн-макет (4 фрейма — два "
        "веб и два мобильных); на второй практике дизайн был переведён в реальную "
        "HTML/CSS-вёрстку и проект был размещён в системе контроля версий Git на "
        "портале GitHub в отдельной ветке студента; на третьей практике приложение "
        "было дополнено серверной логикой, написана документация и записаны "
        "демонстрационные видео.")
    add_paragraph(doc,
        "Все требуемые задачи выполнены, цели практики достигнуты. Полученные навыки "
        "проектирования, командной разработки и работы с системами контроля версий "
        "будут востребованы в дальнейшем обучении и при выполнении выпускной "
        "квалификационной работы.")
    doc.add_page_break()

    # ===== Приложение =====
    add_centered(doc, "ПРИЛОЖЕНИЕ", bold=True, size=16)
    add_paragraph(doc, "Ссылка на GitHub-репозиторий проекта:", bold=True, indent=False)
    add_paragraph(doc, "https://github.com/MatveyHG32/2-PRAKTIKA/tree/magasumov", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Ссылка на дизайн-макет в Figma:", bold=True, indent=False)
    add_paragraph(doc, "https://www.figma.com/file/______________________________ (вставьте сюда ссылку на свой Figma-файл)", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Видеоматериалы:", bold=True, indent=False)
    add_paragraph(doc, "— Видео 1 — структура проекта и листинг кода (приложено отдельным файлом).", indent=False)
    add_paragraph(doc, "— Видео 2 — пользовательская демонстрация работы приложения (приложено отдельным файлом).", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Демо-учётная запись для проверки:", bold=True, indent=False)
    add_paragraph(doc, "Логин: matvey   Пароль: matvey2026", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Запуск проекта на локальной машине:", bold=True, indent=False)
    add_code(doc,
        "git clone -b magasumov https://github.com/MatveyHG32/2-PRAKTIKA\n"
        "cd 2-PRAKTIKA\n"
        "python -m venv .venv\n"
        ".venv/Scripts/activate    # для Windows\n"
        "source .venv/bin/activate  # для Linux/macOS\n"
        "pip install -r requirements.txt\n"
        "python seed_demo.py\n"
        "uvicorn main:app --reload\n"
    )

    out = OUT_DIR / "Отчёт_практика_3.docx"
    doc.save(out)
    print(f"Saved {out}")


if __name__ == "__main__":
    build()
