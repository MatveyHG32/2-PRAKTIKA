"""Generate the Practice 2 report (.docx).

Outputs: docs/reports/Отчёт_практика_2.docx

Структура: титул → индивидуальное задание → введение → 1. Разработка вёрстки
веб-приложения → 2. Размещение проекта на GitHub → 3. Итоги и выводы → 4. Приложения.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "reports"
OUT_DIR.mkdir(parents=True, exist_ok=True)

SCREENS = ROOT / "docs" / "screenshots"
FIGMA = ROOT / "docs" / "figma_mockups"


def add_centered(doc, text, *, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_paragraph(doc, text, *, size=14, indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_heading(doc, text, *, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_image(doc, path: Path, caption: str, width_inches: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(12)
    cap_run.font.name = "Times New Roman"


def add_code(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(11)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    # ===== Титул =====
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", size=12)
    add_centered(doc, "ФГБОУ ВО «Брянский государственный инженерно-технологический университет»", size=12)
    add_centered(doc, "Многопрофильный колледж", size=12)
    add_centered(doc, "Кафедра «Информационные технологии»", size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "ОТЧЁТ", bold=True, size=18)
    add_centered(doc, "по учебной практике по технологии разработки программного обеспечения", size=14)
    add_centered(doc, "(практика 2)", size=14)
    add_centered(doc, "ОП – 02068025-09.02.07-250.25", size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "Студент _________________________________________ Магасумов М. И.", indent=False)
    add_paragraph(doc, "Группа ИСП(спо)-209-2                          № зачётной книжки: 24-2.250", indent=False)
    add_paragraph(doc, "Руководитель от вуза ________________ преподаватель И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Нормоконтроль ______________________ преподаватель И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Доступ к защите «04» 06 2026 г. ___ преподаватель И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Дата защиты «05» 06 2026 г.  Оценка __________________", indent=False)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Брянск 2026", size=12)
    doc.add_page_break()

    # ===== Индивидуальное задание =====
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", size=12)
    add_centered(doc, "Федеральное государственное бюджетное образовательное учреждение", size=12)
    add_centered(doc, "высшего образования", size=12)
    add_centered(doc, "«Брянский государственный инженерно-технологический университет»", size=12)
    add_centered(doc, "Многопрофильный колледж", size=12)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Индивидуальное задание", bold=True, size=16)
    add_centered(doc, "по учебной практике по технологии разработки программного обеспечения", size=14)
    doc.add_paragraph()
    add_paragraph(doc, "Обучающийся: Магасумов М. И.", indent=False)
    add_paragraph(doc, "Руководитель практики от вуза: И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Сроки прохождения практики: 23.05.2026 – 05.06.2026", indent=False)
    add_paragraph(doc, "Место прохождения практики: кафедра ИТ БГИТУ", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Содержание индивидуального задания:", bold=True, indent=False)
    add_paragraph(doc, "ВВЕДЕНИЕ ............................................................................................................. 3", indent=False)
    add_paragraph(doc, "1. Разработка вёрстки веб-приложения ................................................................ 4", indent=False)
    add_paragraph(doc, "2. Размещение проекта на GitHub ........................................................................ 8", indent=False)
    add_paragraph(doc, "3. Итоги и выводы ..................................................................................................10", indent=False)
    add_paragraph(doc, "4. Приложения .......................................................................................................11", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Дата выдачи задания 23 мая 2026 г.", indent=False)
    add_paragraph(doc, "Руководитель практики ________________ И. В. Мартыненко", indent=False)
    add_paragraph(doc, "Задание принял к исполнению ____________", indent=False)
    doc.add_page_break()

    # ===== Введение =====
    add_centered(doc, "ВВЕДЕНИЕ", bold=True, size=16)
    add_paragraph(
        doc,
        "Цель учебной практики 2 — разработать вёрстку веб-приложения «Планировщик задач» "
        "в полном соответствии с дизайн-макетом, выполненным в Figma на практике 1, "
        "а также разместить исходный код проекта в системе контроля версий Git на портале "
        "GitHub в отдельной ветке студента.",
    )
    add_paragraph(
        doc,
        "Актуальность темы по-прежнему обусловлена потребностью в простых и понятных "
        "инструментах для планирования учебных задач. Реализация front-end части "
        "позволяет проверить корректность дизайн-макета на практике: убедиться в "
        "удобстве выбранной цветовой схемы, в адекватности размеров элементов и "
        "в правильности построения адаптивной сетки.",
    )
    add_paragraph(
        doc,
        "В рамках практики были выполнены следующие задачи: написана HTML-структура "
        "страниц, разработаны CSS-правила, реализующие тёмно-фиолетовую цветовую "
        "схему дизайн-макета и адаптивную раскладку, подключён минимальный "
        "JavaScript для обработки модальных окон, а также подключён шаблонизатор "
        "Jinja2 для динамической отрисовки данных из базы данных SQLite.",
    )
    doc.add_page_break()

    # ===== 1. Разработка вёрстки =====
    add_centered(doc, "1. Разработка вёрстки веб-приложения", bold=True, size=16)

    add_heading(doc, "1.1. Технологический стек")
    add_paragraph(
        doc,
        "Для разработки веб-приложения выбран следующий стек: Python 3.12 и "
        "веб-фреймворк FastAPI на стороне сервера, реляционная СУБД SQLite "
        "и ORM SQLAlchemy 2.0 — для хранения данных. На стороне клиента используется "
        "шаблонизатор Jinja2, написаны HTML-страницы по семантической разметке "
        "и подключены каскадные таблицы стилей CSS3 (без сторонних фреймворков). "
        "Минимальный JavaScript обеспечивает работу модальных окон и подтверждение "
        "удаления.",
    )

    add_heading(doc, "1.2. HTML-разметка")
    add_paragraph(
        doc,
        "Все HTML-страницы построены на базе общего шаблона base.html, в котором "
        "вынесены: «шапка» сайта с навигацией, блок входа/выхода, основной "
        "контейнер. Дочерние шаблоны (dashboard.html, task_detail.html, "
        "projects.html, login.html, register.html) определяют только содержимое "
        "блока main. Это уменьшает дублирование разметки и облегчает поддержку.",
    )
    add_paragraph(
        doc,
        "Каждая задача на дашборде представлена «карточкой» с цветной полосой "
        "слева, отражающей приоритет (зелёный — низкий, оранжевый — средний, "
        "красный — высокий), и набором цветных «чипов» сверху (проект и "
        "приоритет). Карточки сгруппированы в три колонки по статусам.",
    )

    add_heading(doc, "1.3. CSS-стили")
    add_paragraph(
        doc,
        "Все стили вынесены в файл static/css/style.css. В нём заданы корневые "
        "CSS-переменные (цвета, отступы, радиусы скругления, тени) — это позволяет "
        "при необходимости легко поменять цветовую схему всего сайта, не правя "
        "конкретные правила. Раскладка строится на CSS Grid и Flexbox.",
    )
    add_code(doc,
        ":root {\n"
        "    --primary: #6366F1;\n"
        "    --primary-dark: #4F46E5;\n"
        "    --bg: #F4F5F9;\n"
        "    --card: #FFFFFF;\n"
        "    --text: #1F2233;\n"
        "    --muted: #6B7080;\n"
        "    --success: #10B981;\n"
        "    --warning: #F59E0B;\n"
        "    --danger:  #EF4444;\n"
        "    --radius: 12px;\n"
        "    --shadow: 0 2px 8px rgba(31, 34, 51, 0.08);\n"
        "}\n"
    )
    add_paragraph(
        doc,
        "Для адаптивности добавлено медиа-правило @media (max-width: 920px): "
        "канбан-доска перестраивается в одну колонку, сетка статистических карточек — "
        "в две, а навигация переходит на новую строку.",
    )

    add_heading(doc, "1.4. JavaScript")
    add_paragraph(
        doc,
        "JavaScript-логика сведена к минимуму. Модальные окна управляются "
        "через data-атрибуты: на кнопке «+ Новая задача» поставлен data-open="
        "\"task-modal\", в самом окне — id=\"task-modal\"; скрипт static/js/app.js "
        "перехватывает клики и переключает CSS-класс .open у элемента.",
    )

    add_heading(doc, "1.5. Результат")
    add_paragraph(
        doc,
        "Готовая вёрстка веб-приложения полностью соответствует дизайн-макету "
        "Figma из практики 1. На рисунках 1–4 показаны те же экраны, что и в "
        "макетах, но уже в виде работающего веб-приложения, отрисованного "
        "браузером Google Chrome.",
    )

    if (SCREENS / "web_01_dashboard.png").exists():
        add_image(doc, SCREENS / "web_01_dashboard.png",
                  "Рисунок 1 — Главный экран (десктоп)")
    if (SCREENS / "web_02_task_detail.png").exists():
        add_image(doc, SCREENS / "web_02_task_detail.png",
                  "Рисунок 2 — Карточка задачи (десктоп)")
    if (SCREENS / "mobile_01_dashboard.png").exists():
        add_image(doc, SCREENS / "mobile_01_dashboard.png",
                  "Рисунок 3 — Главный экран (мобильная версия)", width_inches=3.0)
    if (SCREENS / "mobile_02_task_detail.png").exists():
        add_image(doc, SCREENS / "mobile_02_task_detail.png",
                  "Рисунок 4 — Карточка задачи (мобильная версия)", width_inches=3.0)

    doc.add_page_break()

    # ===== 2. Размещение проекта на GitHub =====
    add_centered(doc, "2. Размещение проекта на GitHub", bold=True, size=16)

    add_heading(doc, "2.1. Создание репозитория и инициализация Git")
    add_paragraph(
        doc,
        "Проект помещён под управление системы контроля версий Git. В корне "
        "проекта инициализирован пустой репозиторий командой git init, добавлен "
        "файл .gitignore для исключения временных файлов (виртуальное окружение, "
        "локальная база данных, кэш Python). Затем все исходные файлы добавлены "
        "в индекс командой git add -A и зафиксированы первым коммитом с подробным "
        "описанием.",
    )

    add_heading(doc, "2.2. Ветка по фамилии студента")
    add_paragraph(
        doc,
        "В соответствии с требованиями преподавателя коммиты сделаны не в ветке "
        "main/master, а в отдельной ветке студента с именем magasumov. Это позволяет "
        "при необходимости легко объединять работу нескольких студентов в общем "
        "репозитории учебной группы. Перевод в новую ветку выполнен командой "
        "git checkout -b magasumov.",
    )

    add_heading(doc, "2.3. Заливка кода на GitHub")
    add_paragraph(
        doc,
        "В личном кабинете на GitHub создан публичный репозиторий task-planner. "
        "Локальный репозиторий привязан к удалённому через git remote add origin "
        "и заливается командой git push -u origin magasumov. Все необходимые "
        "файлы: исходный код приложения, шаблоны, стили, скрипты, документация "
        "и дизайн-макеты — доступны по ссылке в ветке magasumov.",
    )

    add_heading(doc, "2.4. История коммитов")
    add_paragraph(
        doc,
        "В соответствии с требованиями преподавателя коммиты появлялись не перед "
        "сдачей проекта, а в течение всего периода практики. Список коммитов "
        "можно посмотреть командой git log --oneline или в веб-интерфейсе "
        "GitHub на вкладке Commits.",
    )

    doc.add_page_break()

    # ===== 3. Итоги и выводы =====
    add_centered(doc, "3. Итоги и выводы", bold=True, size=16)
    add_paragraph(
        doc,
        "В рамках практики 2 разработана полноценная front-end часть веб-приложения "
        "«Планировщик задач». Все экраны, спроектированные в Figma на практике 1, "
        "реализованы в виде HTML-страниц с CSS-стилями. Дополнительно подключён "
        "шаблонизатор Jinja2, что позволило разделить логику отрисовки и данные.",
    )
    add_paragraph(
        doc,
        "Адаптивная вёрстка проверена в режиме DevTools браузера Google Chrome "
        "(панель Responsive). При ширине окна меньше 920 пикселей канбан-доска "
        "перестраивается в одну колонку, навигация переносится на новую строку, "
        "формы остаются читаемыми и удобными для заполнения.",
    )
    add_paragraph(
        doc,
        "Исходный код проекта размещён на GitHub в отдельной ветке magasumov "
        "и доступен для проверки. В практике 3 планируется реализовать backend-часть "
        "(API эндпоинты, подключение к SQLite, аутентификация), провести функциональное "
        "тестирование и записать демонстрационные видео.",
    )

    doc.add_page_break()

    # ===== 4. Приложения =====
    add_centered(doc, "4. Приложения", bold=True, size=16)
    add_paragraph(doc, "Ссылка на GitHub:", bold=True, indent=False)
    add_paragraph(doc, "https://github.com/MatveyHG32/2-PRAKTIKA/tree/magasumov", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Ссылка на Figma-макет:", bold=True, indent=False)
    add_paragraph(doc, "https://www.figma.com/design/OlEfN3Mh26JJZbXSjgwhTP/Untitled?node-id=0-1&t=1KSjV01RB9biU0kb-1", indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Скриншоты приложения и Figma-фреймы:", bold=True, indent=False)
    add_paragraph(
        doc,
        "Все скриншоты приведены на рисунках 1–4 раздела 1.5. Кроме того, "
        "все исходные PNG-файлы доступны в каталоге docs/screenshots/ "
        "и docs/figma_mockups/ внутри репозитория.",
    )
    doc.add_paragraph()
    add_paragraph(doc, "Запуск проекта на локальной машине:", bold=True, indent=False)
    add_code(doc,
        "git clone -b magasumov https://github.com/MatveyHG32/2-PRAKTIKA\n"
        "cd 2-PRAKTIKA\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install -r requirements.txt\n"
        "python seed_demo.py\n"
        "uvicorn main:app --reload\n"
    )
    add_paragraph(doc, "Демо-логин: matvey / Пароль: matvey2026", indent=False)

    out = OUT_DIR / "Отчёт_практика_2.docx"
    doc.save(out)
    print(f"Saved {out}")


if __name__ == "__main__":
    build()
