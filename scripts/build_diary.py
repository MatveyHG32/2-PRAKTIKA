"""Generate the practice 3 diary (.docx).

Outputs: docs/reports/Дневник_практики.docx

Дневник производственной практики 3, табличный формат:
дата → выполненная работа → отметка руководителя.
"""
from pathlib import Path
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "reports"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def add_centered(doc, text, *, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_para(doc, text, *, indent=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_cell_text(cell, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    set_cell_borders(cell)


# Дни производственной практики 3 и работы по ним.
# Практика 3: 08.06.2026 – 25.06.2026, рабочие дни (пн-пт).
WORK_DAYS = [
    ("08.06.2026", "Инструктаж по технике безопасности. Постановка задачи на практику 3: подготовка демонстрационных видео, оформление дневника и итогового отчёта.", 6),
    ("09.06.2026", "Изучение требований к итоговому отчёту по производственной практике. Анализ образца отчёта.", 6),
    ("10.06.2026", "Подготовка структуры итогового отчёта: разделы «Архитектура приложения», «Основные компоненты», «Алгоритм работы».", 6),
    ("11.06.2026", "Написание раздела 1 — проектирование архитектуры. Описание стека (FastAPI, SQLAlchemy, SQLite, Jinja2).", 6),
    ("15.06.2026", "Написание раздела 2 — основные компоненты приложения (модели данных, роутеры, шаблоны).", 6),
    ("16.06.2026", "Написание раздела 3 — алгоритм работы (схема взаимодействия пользователя и системы).", 6),
    ("17.06.2026", "Описание индивидуального задания: реализация авторизации, CRUD задач, проектов и комментариев.", 6),
    ("18.06.2026", "Проверка вёрстки во всех браузерах, исправление мелких CSS-замечаний.", 6),
    ("19.06.2026", "Подготовка сценария первого видео — обзор структуры проекта и листинг кода.", 6),
    ("22.06.2026", "Запись видео 1: проход по структуре каталогов проекта и постраничное чтение основных файлов кода.", 6),
    ("23.06.2026", "Запись видео 2: демонстрация приложения со стороны пользователя (вход, создание задачи, комментарии, смена статуса, проекты, мобильная версия).", 6),
    ("24.06.2026", "Финальная сборка итогового отчёта, оформление приложений. Заполнение дневника производственной практики.", 6),
    ("25.06.2026", "Подготовка к защите: повторная проверка всех файлов в репозитории, подсчёт коммитов, согласование с руководителем.", 6),
]


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # ===== Титул =====
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", size=12)
    add_centered(doc, "ФГБОУ ВО «Брянский государственный инженерно-технологический университет»", size=12)
    add_centered(doc, "Многопрофильный колледж", size=12)
    add_centered(doc, "Кафедра «Информационные технологии»", size=12)
    for _ in range(4):
        doc.add_paragraph()

    add_centered(doc, "ДНЕВНИК", bold=True, size=22)
    add_centered(doc, "производственной практики", bold=True, size=16)
    add_centered(doc, "по технологии разработки программного обеспечения", size=14)
    add_centered(doc, "(практика 3)", size=14)
    doc.add_paragraph()
    add_centered(doc, "Тема индивидуального задания:", size=12)
    add_centered(doc, "«Планировщик задач» — веб-приложение на FastAPI + SQLite", bold=True, size=12)
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, "Студент:           Магасумов М. И.", size=12)
    add_para(doc, "Группа:            ИСП(спо)-209-2", size=12)
    add_para(doc, "Сроки практики:    08.06.2026 – 25.06.2026", size=12)
    add_para(doc, "Руководитель:      канд. экон. наук, доц. Н. Ю. Азаренко", size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "Брянск 2026", size=12)

    doc.add_page_break()

    # ===== Таблица «Дневник» =====
    add_centered(doc, "ДНЕВНИК прохождения практики", bold=True, size=14)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Шапка
    hdr = table.rows[0].cells
    headers = ["Дата", "Содержание выполненной работы", "Кол-во часов", "Подпись руководителя"]
    for i, text in enumerate(headers):
        set_cell_text(hdr[i], text, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Колоночные ширины
    widths = (Cm(2.4), Cm(10.5), Cm(1.6), Cm(2.5))
    for i, w in enumerate(widths):
        hdr[i].width = w

    # Строки
    total_hours = 0
    for day, work, hours in WORK_DAYS:
        row = table.add_row().cells
        set_cell_text(row[0], day, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], work, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[2], str(hours), size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[3], "", size=11)
        for i, w in enumerate(widths):
            row[i].width = w
        total_hours += hours

    # Итого
    total_row = table.add_row().cells
    set_cell_text(total_row[0], "Итого:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(total_row[1], "", size=11)
    set_cell_text(total_row[2], str(total_hours), bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(total_row[3], "", size=11)

    doc.add_paragraph()
    add_para(doc, "Руководитель практики от вуза _______________ Н. Ю. Азаренко", size=12)
    add_para(doc, "Студент _______________ Магасумов М. И.", size=12)
    doc.add_paragraph()
    add_para(doc, "Дата заполнения: 25 июня 2026 г.", size=12)

    out = OUT_DIR / "Дневник_практики.docx"
    doc.save(out)
    print(f"Saved {out}")


if __name__ == "__main__":
    build()
